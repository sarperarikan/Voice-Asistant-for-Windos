# -*- coding: utf-8 -*-
import os
import json
import threading
import webview
import keyboard
from http.server import SimpleHTTPRequestHandler, HTTPServer
from docx import Document
from Levenshtein import distance as levenshtein_distance
import logging

# Logging yapılandırması
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class Api:
    def __init__(self, command_executor):
        self.command_executor = command_executor

    def send_result(self, result):
        """Handle the received result from the voice recognition."""
        result = self.normalize_command(result)
        logging.info(f"Received result: {result}")
        self.command_executor.execute_command(result)

    def save_to_word(self, content):
        """Save the given content to a Word document in the user's Documents folder."""
        try:
            document = Document()
            document.add_paragraph(content)
            documents_path = os.path.join(os.path.expanduser('~'), 'Documents')
            file_path = os.path.join(documents_path, 'voice-assistant.docx')
            document.save(file_path)
            return "Word belgesi başarıyla Belgelerim klasörüne kaydedildi."
        except Exception as e:
            logging.error(f"Word belgesi kaydedilirken hata oluştu: {str(e)}")
            return f"Word belgesi kaydedilirken hata oluştu: {str(e)}"

    def normalize_command(self, command):
        """Normalize the command string, considering Turkish characters."""
        command = command.lower().strip()
        command = command.replace('ç', 'c').replace('ğ', 'g').replace('ı', 'i')
        command = command.replace('ö', 'o').replace('ş', 's').replace('ü', 'u')
        return ' '.join(command.split())

class KeyboardCommands:
    def __init__(self):
        self.commands = self.load_commands('keyboard.json')
        self.apps = self.load_commands('apps.json')

    def load_commands(self, filename):
        """Load commands from a JSON file."""
        try:
            with open(filename, 'r', encoding='utf-8') as file:
                commands = json.load(file)
                logging.info(f"Commands from {filename} loaded successfully.")
                return {self.normalize_command(k): v for k, v in commands.items()}
        except FileNotFoundError:
            logging.error(f"{filename} file not found.")
            return {}
        except json.JSONDecodeError:
            logging.error(f"Error decoding {filename} file.")
            return {}

    def execute_command(self, command_name):
        """Execute the command that matches the given command name."""
        command_name = self.normalize_command(command_name)
        logging.info(f"Trying to match command: {command_name}")
        
        best_match = self.find_best_match(command_name, self.commands)
        if best_match:
            command = self.commands[best_match]
            logging.info(f"Matched command: {command_name} -> {command}")
            self._execute_command(command, use_os=command in self.apps)
        else:
            logging.warning(f"Command not found: {command_name}")

    def find_best_match(self, command_name, commands):
        """Find the best match for a command using Levenshtein distance."""
        best_match = None
        min_distance = float('inf')
        for key in commands.keys():
            distance = levenshtein_distance(command_name, key)
            if distance < min_distance:
                min_distance = distance
                best_match = key
        logging.info(f"Best match: {best_match} with distance {min_distance}")
        return best_match if min_distance <= 2 else None  # Allowing a maximum distance of 2 for similar commands

    def normalize_command(self, command):
        """Normalize the command string, considering Turkish characters."""
        command = command.lower().strip()
        command = command.replace('ç', 'c').replace('ğ', 'g').replace('ı', 'i')
        command = command.replace('ö', 'o').replace('ş', 's').replace('ü', 'u')
        return ' '.join(command.split())

    def _execute_command(self, command, use_os=False):
        """Execute the given command."""
        logging.info(f"Executing command: {command}")
        if use_os:
            os.system(command)
        else:
            keyboard.press_and_release(command)

class CustomHTTPRequestHandler(SimpleHTTPRequestHandler):
    def do_GET(self):
        """Handle GET requests."""
        if self.path == '/keyboard.json':
            self._send_json_response('keyboard.json')
        elif self.path == '/apps.json':
            self._send_json_response('apps.json')
        else:
            super().do_GET()

    def _send_json_response(self, filename):
        """Send a JSON response from a file."""
        try:
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            with open(filename, 'r', encoding='utf-8') as file:
                self.wfile.write(file.read().encode('utf-8'))
        except FileNotFoundError:
            self.send_error(404, f"File not found: {filename}")
        except Exception as e:
            self.send_error(500, f"Internal server error: {str(e)}")

    def do_POST(self):
        """Handle POST requests."""
        if self.path == '/save_to_word':
            self._handle_save_to_word()
        elif self.path == '/execute_command':
            self._handle_execute_command()
        elif self.path == '/add_command':
            self._handle_add_command()

    def _handle_save_to_word(self):
        """Handle the save to word request."""
        content_length = int(self.headers['Content-Length'])
        post_data = self.rfile.read(content_length)
        data = json.loads(post_data)
        content = data.get('content', '')

        response = {}
        try:
            document = Document()
            document.add_paragraph(content)
            documents_path = os.path.join(os.path.expanduser('~'), 'Documents')
            file_path = os.path.join(documents_path, 'voice-assistant.docx')
            document.save(file_path)
            response['success'] = True
            response['message'] = 'Word belgesi başarıyla Belgelerim klasörüne kaydedildi.'
        except Exception as e:
            response['success'] = False
            response['message'] = f'Word belgesi kaydedilirken hata oluştu: {str(e)}'

        self.send_response(200 if response['success'] else 500)
        self.send_header('Content-type', 'application/json')
        self.end_headers()
        self.wfile.write(json.dumps(response).encode('utf-8'))

    def _handle_execute_command(self):
        """Handle the execute command request."""
        content_length = int(self.headers['Content-Length'])
        post_data = self.rfile.read(content_length)
        data = json.loads(post_data)
        command_name = data.get('command', '')

        response = {}
        try:
            self.server.keyboard_commands.execute_command(command_name)
            response['success'] = True
            response['message'] = 'Komut başarıyla yürütüldü.'
        except Exception as e:
            logging.error(f"Error executing command: {str(e)}")
            response['success'] = False
            response['message'] = f'Komut yürütülürken hata oluştu: {str(e)}'

        self.send_response(200 if response['success'] else 500)
        self.send_header('Content-type', 'application/json')
        self.end_headers()
        self.wfile.write(json.dumps(response).encode('utf-8'))

    def _handle_add_command(self):
        """Handle the add command request."""
        content_length = int(self.headers['Content-Length'])
        post_data = self.rfile.read(content_length)
        data = json.loads(post_data)
        command = data.get('command', '')
        action = data.get('action', '')

        response = {}
        try:
            commands = self.server.keyboard_commands.load_commands('keyboard.json')
            commands[command] = action
            with open('keyboard.json', 'w', encoding='utf-8') as file:
                json.dump(commands, file, ensure_ascii=False, indent=4)
            response['success'] = True
            response['message'] = 'Komut başarıyla eklendi.'
        except Exception as e:
            logging.error(f"Error adding command: {str(e)}")
            response['success'] = False
            response['message'] = f'Komut eklenirken hata oluştu: {str(e)}'

        self.send_response(200 if response['success'] else 500)
        self.send_header('Content-type', 'application/json')
        self.end_headers()
        self.wfile.write(json.dumps(response).encode('utf-8'))

def start_server():
    """Start the HTTP server."""
    web_dir = os.path.join(os.path.dirname(__file__), 'resources')
    if (os.path.exists(web_dir)):
        os.chdir(web_dir)
    else:
        logging.error(f"Directory {web_dir} does not exist.")
        return
    server_address = ('', 8000)
    httpd = HTTPServer(server_address, CustomHTTPRequestHandler)
    httpd.keyboard_commands = KeyboardCommands()  # Ekleme: Bu satırla keyboard_commands özelliği ekleniyor
    httpd.serve_forever()

def start_webview():
    """Start the web view with the voice assistant application."""
    command_executor = KeyboardCommands()
    api = Api(command_executor)
    threading.Thread(target=start_server, daemon=True).start()
    webview.create_window('Voice Assistant for Windows', 'http://localhost:8000/index.html', js_api=api)
    webview.start()

if __name__ == '__main__':
    start_webview()
